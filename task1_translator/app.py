import streamlit as st
from deep_translator import GoogleTranslator
from gtts import gTTS
import io
import urllib.parse
from fpdf import FPDF
from docx import Document
import os

st.set_page_config(page_title="SupportBot | Hackathon 2026", page_icon="🌍")

# --- THEME ENGINE ---
def get_current_theme():
    config_dir = ".streamlit"
    config_path = os.path.join(config_dir, "config.toml")
    if not os.path.exists(config_path):
        return "System Default"
    try:
        with open(config_path, "r") as f:
            content = f.read()
            if 'backgroundColor="#0B1D3A"' in content: return "Midnight Blue"
            elif 'backgroundColor="#1a3300"' in content: return "Forest Green"
            elif 'backgroundColor="#0e1117"' in content: return "Dark Mode"
            elif 'backgroundColor="#ffffff"' in content: return "Light Mode"
    except:
        pass
    return "System Default"

def apply_theme(theme_name):
    config_dir = ".streamlit"
    config_path = os.path.join(config_dir, "config.toml")
    
    if not os.path.exists(config_dir):
        os.makedirs(config_dir)
        
    themes = {
        "Light Mode": {"base": "light", "primaryColor": "#ff4b4b", "backgroundColor": "#ffffff", "secondaryBackgroundColor": "#f0f2f6", "textColor": "#000000"},
        "Dark Mode": {"base": "dark", "primaryColor": "#ff4b4b", "backgroundColor": "#0e1117", "secondaryBackgroundColor": "#262730", "textColor": "#ffffff"},
        "Midnight Blue": {"base": "dark", "primaryColor": "#4A90E2", "backgroundColor": "#0B1D3A", "secondaryBackgroundColor": "#071326", "textColor": "#E8F1F2"},
        "Forest Green": {"base": "dark", "primaryColor": "#80ff00", "backgroundColor": "#1a3300", "secondaryBackgroundColor": "#0d1a00", "textColor": "#e6ffcc"},
    }
    
    if theme_name == "System Default":
        if os.path.exists(config_path):
            os.remove(config_path)
            st.rerun()
    elif theme_name in themes:
        t = themes[theme_name]
        config_content = f"""[theme]
base="{t['base']}"
primaryColor="{t['primaryColor']}"
backgroundColor="{t['backgroundColor']}"
secondaryBackgroundColor="{t['secondaryBackgroundColor']}"
textColor="{t['textColor']}"
"""
        current_content = ""
        if os.path.exists(config_path):
            with open(config_path, "r") as f:
                current_content = f.read()
                
        if current_content.strip() != config_content.strip():
            with open(config_path, "w") as f:
                f.write(config_content)
            st.rerun()

# Customization Sidebar
st.sidebar.title("Customization")

current_theme = get_current_theme()
theme_options = ["System Default", "Light Mode", "Dark Mode", "Midnight Blue", "Forest Green"]
selected_theme = st.sidebar.selectbox("App Theme", theme_options, index=theme_options.index(current_theme))

if selected_theme != current_theme:
    apply_theme(selected_theme)

font_size = st.sidebar.slider("Text Font Size", 12, 40, 16)
font_family = st.sidebar.selectbox("Font Style", ["sans-serif", "serif", "monospace", "cursive", "Arial", "Courier New"])

# Apply Font Sizes and Families via CSS (Colors are now handled by Streamlit natively above)
custom_css = f"""
<style>
    textarea {{
        font-size: {font_size}px !important;
    }}
    /* Apply Font Family to major text elements */
    p, h1, h2, h3, h4, h5, h6, li, label, .stMarkdown, .stText, textarea, input, select, button {{
        font-family: {font_family} !important;
    }}
    /* Protect Streamlit Material Icons from being overwritten */
    .material-symbols-rounded, [data-testid="stIconMaterial"], [class*="icon"] {{
        font-family: "Material Symbols Rounded" !important;
    }}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)
# --------------------

st.title("LinguaBridge: Universal Translator")
st.markdown("A fast, free, and robust translation engine built for seamless cross-border communication.")

with st.sidebar.expander("About LinguaBridge"):
    st.markdown("""
    **LinguaBridge** was developed as a comprehensive solution for the hackathon. 
    It leverages high-speed translation APIs to deliver accurate results across 100+ languages without any subscription fees.
    
    *Version 1.0.0*  
    *Built with Python & Streamlit*
    """)

# Get all supported languages
try:
    langs_dict = GoogleTranslator().get_supported_languages(as_dict=True)
except:
    langs_dict = {"english": "en", "urdu": "ur", "arabic": "ar", "spanish": "es", "french": "fr"}
languages = list(langs_dict.keys())

# Right-to-Left Language Check
rtl_languages = ["urdu", "arabic", "persian", "hebrew", "pashto", "sindhi"]

def is_rtl(lang_name):
    return lang_name.lower() in rtl_languages

col1, col2 = st.columns(2)

with col1:
    source_lang = st.selectbox("Source Language", ["auto"] + languages, index=0, format_func=lambda x: x.title(), help="Select the language you are translating from. Use 'Auto' to detect automatically.")
    
    # Dynamic RTL for Source Text Area
    if is_rtl(source_lang):
        st.markdown('<style>div[data-testid="column"]:nth-of-type(1) textarea { direction: rtl !important; text-align: right !important; }</style>', unsafe_allow_html=True)
    else:
        st.markdown('<style>div[data-testid="column"]:nth-of-type(1) textarea { direction: ltr !important; text-align: left !important; }</style>', unsafe_allow_html=True)

    source_text = st.text_area("Source Text", placeholder="Type or paste your text here...", height=200, max_chars=5000)
    
    # Text-to-Speech for Source
    if st.button("Play Source Audio", key="play_src") and source_text.strip():
        with st.spinner("Generating Audio..."):
            try:
                tts_lang = langs_dict.get(source_lang, 'en') if source_lang != "auto" else 'en'
                tts = gTTS(text=source_text, lang=tts_lang)
                fp = io.BytesIO()
                tts.write_to_fp(fp)
                st.audio(fp, format="audio/mp3")
            except Exception as e:
                st.error(f"TTS Error: {e}")

with col2:
    default_target_idx = languages.index("urdu") if "urdu" in languages else 0
    target_lang = st.selectbox("Target Language", languages, index=default_target_idx, format_func=lambda x: x.title())
    
    # Dynamic RTL for Target Text Area
    if is_rtl(target_lang):
        st.markdown('<style>div[data-testid="column"]:nth-of-type(2) textarea { direction: rtl !important; text-align: right !important; }</style>', unsafe_allow_html=True)
    else:
        st.markdown('<style>div[data-testid="column"]:nth-of-type(2) textarea { direction: ltr !important; text-align: left !important; }</style>', unsafe_allow_html=True)
    
    if st.button("Translate", type="primary"):
        if source_text.strip():
            with st.spinner("Translating..."):
                try:
                    source_code = "auto" if source_lang == "auto" else langs_dict[source_lang]
                    target_code = langs_dict[target_lang]
                    
                    translator = GoogleTranslator(source=source_code, target=target_code)
                    st.session_state.translated_text = translator.translate(source_text)
                except Exception as e:
                    st.error(f"Error during translation: {e}")
        else:
            st.warning("Please enter some text to translate.")
            
    display_text = st.session_state.get('translated_text', '')
    
    st.text_area("Translated Text:", value=display_text, height=200)
    
    if display_text:
        # TTS for Target
        if st.button("Play Translated Audio", key="play_tgt"):
            with st.spinner("Generating Audio..."):
                try:
                    tts_lang = langs_dict.get(target_lang, 'en')
                    tts = gTTS(text=display_text, lang=tts_lang)
                    fp = io.BytesIO()
                    tts.write_to_fp(fp)
                    st.audio(fp, format="audio/mp3")
                except Exception as e:
                    st.error(f"TTS Error: {e}")

        # Custom Copy Button
        copy_html = f"""
        <script>
        function copyText() {{
            navigator.clipboard.writeText(`{display_text}`);
            let btn = document.getElementById("copyBtn");
            btn.innerText = "✓ Copied!";
            btn.style.backgroundColor = "#28a745";
            setTimeout(() => {{
                btn.innerText = "📋 Copy Text";
                btn.style.backgroundColor = "#ff4b4b";
            }}, 2000);
        }}
        </script>
        <button id="copyBtn" onclick="copyText()" style="background-color: #ff4b4b; color: white; border: none; padding: 8px 16px; border-radius: 4px; cursor: pointer; font-weight: bold; font-family: sans-serif; width: 100%;">📋 Copy Text</button>
        """
        import streamlit.components.v1 as components
        components.html(copy_html, height=50)

st.divider()

if st.session_state.get('translated_text', ''):
    st.markdown("### Export & Share")
    display_text = st.session_state.translated_text
    
    col_dl1, col_dl2, col_dl3 = st.columns(3)
    
    with col_dl1:
        st.download_button("Download TXT", data=display_text, file_name="translation.txt", mime="text/plain", use_container_width=True)
        
    with col_dl2:
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("helvetica", size=12)
            pdf.set_auto_page_break(auto=True, margin=15)
            # Encoding to latin-1 and replacing unrepresentable chars to prevent FPDF crash
            safe_text = display_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, txt=safe_text)
            pdf_bytes = bytes(pdf.output())
            st.download_button("Download PDF", data=pdf_bytes, file_name="translation.pdf", mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.error(f"PDF Export Error (Try DOCX for Unicode): {e}")
            
    with col_dl3:
        try:
            doc = Document()
            doc.add_heading('LinguaBridge Translation', 0)
            doc.add_paragraph(display_text)
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            st.download_button("Download DOCX", data=doc_stream.getvalue(), file_name="translation.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        except Exception as e:
            st.error(f"DOCX error: {e}")

    st.markdown("**Share to Social Media:**")
    encoded_text = urllib.parse.quote(display_text)
    
    whatsapp_url = f"https://api.whatsapp.com/send?text={encoded_text}"
    twitter_url = f"https://twitter.com/intent/tweet?text={encoded_text}"
    email_url = f"mailto:?subject=Translation&body={encoded_text}"
    
    st.markdown(f"""
    <div style="display: flex; gap: 10px; margin-top: 10px;">
        <a href="{whatsapp_url}" target="_blank" style="text-decoration:none;"><button style="background-color:#25D366; color:white; border:none; padding:10px 20px; border-radius:5px; cursor:pointer; font-weight:bold;">WhatsApp</button></a>
        <a href="{twitter_url}" target="_blank" style="text-decoration:none;"><button style="background-color:#1DA1F2; color:white; border:none; padding:10px 20px; border-radius:5px; cursor:pointer; font-weight:bold;">Twitter</button></a>
        <a href="{email_url}" target="_blank" style="text-decoration:none;"><button style="background-color:#D44638; color:white; border:none; padding:10px 20px; border-radius:5px; cursor:pointer; font-weight:bold;">Email</button></a>
    </div>
    """, unsafe_allow_html=True)
